"""Resume / LinkedIn file → plain text.

Supported extensions: ``pdf``, ``docx``, ``txt``, ``md``, ``html``, ``htm``.
Any other extension returns a :class:`ParsedFile` with ``error`` set so the
UI can show a friendly message instead of crashing.

Imports are lazy: the heavy parsers (``pypdf``, ``python-docx``) load only
when actually needed, which keeps app startup snappy.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


SUPPORTED_EXTENSIONS = ("pdf", "docx", "txt", "md", "html", "htm")


@dataclass
class ParsedFile:
    path: str
    name: str
    ext: str
    size_bytes: int
    text: str
    error: Optional[str]

    @property
    def ok(self) -> bool:
        return bool(self.text) and not self.error


def _format_size(num_bytes: int) -> str:
    if num_bytes >= 1024 * 1024:
        return f"{num_bytes / (1024 * 1024):.1f} MB"
    if num_bytes >= 1024:
        return f"{num_bytes / 1024:.0f} kB"
    return f"{num_bytes} B"


def human_size(num_bytes: int) -> str:
    return _format_size(num_bytes)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "pypdf not installed - run build_exe.bat or 'pip install pypdf'"
        ) from exc
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(chunks)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed - run build_exe.bat or 'pip install python-docx'"
        ) from exc
    doc = Document(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            chunks.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text)
            if row_text:
                chunks.append(row_text)
    return "\n".join(chunks)


def _read_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        from bs4 import BeautifulSoup  # type: ignore[import-not-found]
    except ImportError:
        text = re.sub(r"<[^>]+>", " ", raw)
        return re.sub(r"\s+", " ", text).strip()
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n").strip()


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def parse_file(path_str: str) -> ParsedFile:
    if not path_str:
        return ParsedFile(path="", name="", ext="", size_bytes=0, text="", error="Empty path.")

    path = Path(path_str)
    if not path.is_file():
        return ParsedFile(
            path=str(path),
            name=path.name,
            ext="",
            size_bytes=0,
            text="",
            error="File not found.",
        )

    ext = path.suffix.lower().lstrip(".")
    try:
        size_bytes = path.stat().st_size
    except OSError:
        size_bytes = 0

    if ext not in SUPPORTED_EXTENSIONS:
        return ParsedFile(
            path=str(path),
            name=path.name,
            ext=ext,
            size_bytes=size_bytes,
            text="",
            error=f"Unsupported extension .{ext}. Use PDF, DOCX, TXT or HTML.",
        )

    try:
        if ext == "pdf":
            text = _read_pdf(path)
        elif ext == "docx":
            text = _read_docx(path)
        elif ext in ("html", "htm"):
            text = _read_html(path)
        else:
            text = _read_text(path)
    except Exception as exc:
        return ParsedFile(
            path=str(path),
            name=path.name,
            ext=ext,
            size_bytes=size_bytes,
            text="",
            error=f"Could not parse file: {exc}",
        )

    text = text.strip()
    if not text:
        return ParsedFile(
            path=str(path),
            name=path.name,
            ext=ext,
            size_bytes=size_bytes,
            text="",
            error="The file looks empty. If it is a scanned PDF, OCR is needed first.",
        )

    return ParsedFile(
        path=str(path),
        name=path.name,
        ext=ext,
        size_bytes=size_bytes,
        text=text,
        error=None,
    )


def is_supported(filename: str) -> bool:
    ext = Path(filename).suffix.lower().lstrip(".")
    return ext in SUPPORTED_EXTENSIONS
