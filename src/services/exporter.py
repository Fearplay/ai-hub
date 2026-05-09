"""Markdown → MD / HTML / DOCX / PDF writers.

The renderer is intentionally narrow: it understands the subset of
markdown that the AI Career pipeline produces (headings up to H3, bold,
italics, ordered + unordered bullet lists, paragraphs, soft line breaks).
Everything else is rendered as plain paragraph text.

ATS friendliness drives all formatting decisions:

* single column
* no tables
* readable serif body / sans heading
* generous left margin
* page breaks on H1 only

PDF uses :mod:`reportlab` (pure Python - no GTK / Office / Pandoc).
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^([*\-+])\s+(.*)$")
_ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
_PDF_FONT_CACHE: tuple[str, str] | None = None


def _resolve_pdf_fonts() -> tuple[str, str]:
    """Pick a Unicode-safe font pair for ReportLab PDFs.

    Built-in Helvetica lacks many Central-European glyphs, which produces
    black squares for Czech text. We prefer common system TTF fonts and
    gracefully fall back to Helvetica when none are available.
    """
    global _PDF_FONT_CACHE
    if _PDF_FONT_CACHE is not None:
        return _PDF_FONT_CACHE

    fallback = ("Helvetica", "Helvetica-Bold")
    try:
        from reportlab.pdfbase import pdfmetrics  # type: ignore[import-not-found]
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore[import-not-found]
    except ImportError:
        _PDF_FONT_CACHE = fallback
        return fallback

    candidates: list[tuple[Path, Path]] = []
    if os.name == "nt":
        fonts_dir = Path(os.environ.get("WINDIR", "C:\\Windows")) / "Fonts"
        candidates.extend(
            [
                (fonts_dir / "arial.ttf", fonts_dir / "arialbd.ttf"),
                (fonts_dir / "calibri.ttf", fonts_dir / "calibrib.ttf"),
                (fonts_dir / "segoeui.ttf", fonts_dir / "segoeuib.ttf"),
            ]
        )
    elif os.name == "posix":
        candidates.extend(
            [
                (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
                (Path("/Library/Fonts/Arial.ttf"), Path("/Library/Fonts/Arial Bold.ttf")),
            ]
        )

    for normal_path, bold_path in candidates:
        if not normal_path.exists():
            continue
        try:
            if "AIHubUnicodeBody" not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont("AIHubUnicodeBody", str(normal_path)))
            if bold_path.exists():
                if "AIHubUnicodeBold" not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont("AIHubUnicodeBold", str(bold_path)))
                _PDF_FONT_CACHE = ("AIHubUnicodeBody", "AIHubUnicodeBold")
            else:
                _PDF_FONT_CACHE = ("AIHubUnicodeBody", "AIHubUnicodeBody")
            return _PDF_FONT_CACHE
        except Exception:
            continue

    _PDF_FONT_CACHE = fallback
    return fallback


@dataclass
class Block:
    kind: str  # "h1" / "h2" / "h3" / "p" / "ul" / "ol"
    text: str = ""
    items: list[str] | None = None


def parse_markdown(text: str) -> list[Block]:
    """Lightweight markdown parser tuned for our outputs."""
    blocks: list[Block] = []
    if not text:
        return blocks

    lines = text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(3, len(heading.group(1)))
            blocks.append(Block(kind=f"h{level}", text=heading.group(2).strip()))
            i += 1
            continue

        bullet = _BULLET_RE.match(stripped)
        if bullet:
            items: list[str] = [bullet.group(2).strip()]
            i += 1
            while i < len(lines):
                nxt = _BULLET_RE.match(lines[i].strip())
                if not nxt or not lines[i].strip():
                    break
                items.append(nxt.group(2).strip())
                i += 1
            blocks.append(Block(kind="ul", items=items))
            continue

        ordered = _ORDERED_RE.match(stripped)
        if ordered:
            items = [ordered.group(2).strip()]
            i += 1
            while i < len(lines):
                nxt = _ORDERED_RE.match(lines[i].strip())
                if not nxt or not lines[i].strip():
                    break
                items.append(nxt.group(2).strip())
                i += 1
            blocks.append(Block(kind="ol", items=items))
            continue

        para_lines: list[str] = [stripped]
        i += 1
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                break
            if _HEADING_RE.match(line) or _BULLET_RE.match(line) or _ORDERED_RE.match(line):
                break
            para_lines.append(line)
            i += 1
        blocks.append(Block(kind="p", text=" ".join(para_lines)))
    return blocks


def export_markdown(text: str, target: Path) -> Path:
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(text or "", encoding="utf-8")
    return target


# Shared "play nicely with PDF print + text selection" rules.
# * ``print-color-adjust:exact`` keeps backgrounds + accent colours
#   when Chrome / Edge print the page (without it the headings lose
#   their colour and the page looks washed out).
# * ``::selection`` gives the user a high-contrast highlight when they
#   drag-select inside the PDF preview - prior CSS left the highlight
#   colour inheriting from the parent which on dark accent banners
#   produced "I can't read what I just selected" complaints.
# * ``a { ... text-decoration:underline }`` makes links visibly
#   clickable in Chrome's PDF viewer.
_PRINT_BASE_CSS = (
    "  *{-webkit-print-color-adjust:exact;print-color-adjust:exact;}\n"
    "  ::selection{background:#FDE68A;color:#0F172A;}\n"
    "  a{color:#1D4ED8;text-decoration:underline;}\n"
    "  a:visited{color:#1D4ED8;}\n"
)

_ATS_CSS = (
    _PRINT_BASE_CSS
    + "  body { font-family: 'Segoe UI', Arial, sans-serif; color: #1f2937; "
    "max-width: 760px; margin: 32px auto; padding: 0 24px; line-height: 1.5; }\n"
    "  h1 { font-size: 24px; margin: 18px 0 8px; }\n"
    "  h2 { font-size: 18px; margin: 16px 0 6px; }\n"
    "  h3 { font-size: 15px; margin: 12px 0 4px; }\n"
    "  p, li { font-size: 13px; }\n"
    "  ul, ol { margin: 4px 0 12px 22px; }\n"
)

# "Modern" style is for the HR-eye CV / cover letter exports. Same single-
# column layout (so the document still parses cleanly when a recruiter
# pastes it back into a tool) but with a coloured H1, accent rules, and
# a more confident type stack.
_MODERN_CSS = (
    _PRINT_BASE_CSS
    + "  :root { --accent: #6366F1; --ink: #0F172A; --muted: #475569; }\n"
    "  body { font-family: 'Inter', 'Segoe UI', 'Helvetica Neue', Arial, "
    "sans-serif; color: var(--ink); max-width: 780px; margin: 40px auto; "
    "padding: 0 36px; line-height: 1.55; }\n"
    "  h1 { font-size: 30px; margin: 6px 0 4px; color: var(--accent); "
    "letter-spacing: -0.5px; }\n"
    "  h1 + p, h1 + em { color: var(--muted); font-size: 14px; "
    "margin-top: 0; }\n"
    "  h2 { font-size: 18px; margin: 28px 0 8px; padding-bottom: 4px; "
    "border-bottom: 2px solid var(--accent); display: inline-block; }\n"
    "  h3 { font-size: 15px; margin: 16px 0 4px; color: var(--ink); }\n"
    "  hr { border: 0; border-top: 1px solid #E2E8F0; margin: 22px 0; }\n"
    "  p, li { font-size: 13.5px; color: var(--ink); }\n"
    "  em { color: var(--muted); }\n"
    "  ul, ol { margin: 6px 0 14px 22px; }\n"
    "  ul li { margin-bottom: 4px; }\n"
    "  strong { color: var(--ink); }\n"
    "  a{color:var(--accent);}\n"
)


def export_html(
    text: str,
    target: Path,
    *,
    title: str = "Document",
    style: str = "ats",
) -> Path:
    try:
        import markdown2  # type: ignore[import-not-found]

        body = markdown2.markdown(text or "", extras=["fenced-code-blocks", "tables", "strike"])
    except ImportError:
        body = _blocks_to_html(parse_markdown(text or ""))

    css = _MODERN_CSS if style == "modern" else _ATS_CSS
    html = (
        "<!doctype html>\n"
        "<html lang=\"en\">\n<head>\n"
        "<meta charset=\"utf-8\" />\n"
        f"<title>{_escape_html(title)}</title>\n"
        "<style>\n"
        + css
        + "</style>\n"
        "</head>\n<body>\n"
        + body
        + "\n</body>\n</html>\n"
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(html, encoding="utf-8")
    return target


def export_docx(text: str, target: Path, *, title: str = "Document") -> Path:
    try:
        from docx import Document  # type: ignore[import-not-found]
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx not installed - run pip install -r requirements.txt") from exc

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    blocks = parse_markdown(text or "")
    for block in blocks:
        if block.kind in ("h1", "h2", "h3"):
            level = int(block.kind[1])
            doc.add_heading(block.text, level=level)
        elif block.kind == "p":
            doc.add_paragraph(_strip_inline(block.text))
        elif block.kind == "ul":
            for item in block.items or []:
                doc.add_paragraph(_strip_inline(item), style="List Bullet")
        elif block.kind == "ol":
            for item in block.items or []:
                doc.add_paragraph(_strip_inline(item), style="List Number")

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    return target


def export_pdf(
    text: str,
    target: Path,
    *,
    title: str = "Document",
    style: str = "ats",
) -> Path:
    try:
        from reportlab.lib.colors import HexColor  # type: ignore[import-not-found]
        from reportlab.lib.pagesizes import A4  # type: ignore[import-not-found]
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-not-found]
        from reportlab.lib.units import cm  # type: ignore[import-not-found]
        from reportlab.platypus import (  # type: ignore[import-not-found]
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError as exc:
        raise RuntimeError("reportlab not installed - run pip install -r requirements.txt") from exc

    body_font, heading_font = _resolve_pdf_fonts()
    styles = getSampleStyleSheet()
    if style == "modern":
        accent = HexColor("#6366F1")
        ink = HexColor("#0F172A")
        body = ParagraphStyle(
            "body",
            parent=styles["BodyText"],
            fontName=body_font,
            fontSize=10.5,
            leading=15,
            spaceAfter=5,
            textColor=ink,
        )
        h1 = ParagraphStyle(
            "h1",
            parent=styles["Heading1"],
            fontName=heading_font,
            fontSize=22,
            leading=26,
            spaceAfter=4,
            textColor=accent,
        )
        h2 = ParagraphStyle(
            "h2",
            parent=styles["Heading2"],
            fontName=heading_font,
            fontSize=14,
            leading=18,
            spaceBefore=14,
            spaceAfter=6,
            textColor=ink,
            borderColor=accent,
            borderPadding=(0, 0, 4, 0),
            borderWidth=0,
        )
        h3 = ParagraphStyle(
            "h3",
            parent=styles["Heading3"],
            fontName=heading_font,
            fontSize=12,
            leading=15,
            spaceAfter=3,
            textColor=ink,
        )
    else:
        body = ParagraphStyle(
            "body",
            parent=styles["BodyText"],
            fontName=body_font,
            fontSize=10,
            leading=14,
            spaceAfter=4,
        )
        h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=heading_font, fontSize=18, leading=22, spaceAfter=8)
        h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=heading_font, fontSize=14, leading=18, spaceAfter=6)
        h3 = ParagraphStyle("h3", parent=styles["Heading3"], fontName=heading_font, fontSize=11, leading=14, spaceAfter=4)

    doc = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title=title,
    )

    flowables: list = []
    blocks = parse_markdown(text or "")
    for block in blocks:
        if block.kind == "h1":
            flowables.append(Paragraph(_inline_to_html(block.text), h1))
        elif block.kind == "h2":
            flowables.append(Paragraph(_inline_to_html(block.text), h2))
        elif block.kind == "h3":
            flowables.append(Paragraph(_inline_to_html(block.text), h3))
        elif block.kind == "p":
            flowables.append(Paragraph(_inline_to_html(block.text), body))
            flowables.append(Spacer(1, 4))
        elif block.kind == "ul":
            items = [
                ListItem(Paragraph(_inline_to_html(item), body), leftIndent=12)
                for item in (block.items or [])
            ]
            if items:
                flowables.append(ListFlowable(items, bulletType="bullet", leftIndent=18))
                flowables.append(Spacer(1, 4))
        elif block.kind == "ol":
            items = [
                ListItem(Paragraph(_inline_to_html(item), body), leftIndent=12)
                for item in (block.items or [])
            ]
            if items:
                flowables.append(ListFlowable(items, bulletType="1", leftIndent=18))
                flowables.append(Spacer(1, 4))

    if not flowables:
        flowables.append(Paragraph(_escape_html(text or ""), body))

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.build(flowables)
    return target


def export_all(text: str, folder: Path, basename: str, *, title: Optional[str] = None) -> dict[str, Path]:
    """Convenience helper - writes MD, HTML, DOCX, PDF using one base name."""
    folder.mkdir(parents=True, exist_ok=True)
    title = title or basename
    out: dict[str, Path] = {}
    out["md"] = export_markdown(text, folder / f"{basename}.md")
    out["html"] = export_html(text, folder / f"{basename}.html", title=title)
    try:
        out["docx"] = export_docx(text, folder / f"{basename}.docx", title=title)
    except RuntimeError:
        pass
    try:
        out["pdf"] = export_pdf(text, folder / f"{basename}.pdf", title=title)
    except RuntimeError:
        pass
    return out


def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


_LINK_COLOR = "#1D4ED8"


def _inline_to_html(text: str) -> str:
    """Inline markdown -> ReportLab Paragraph mini-XML, with clickable links.

    ReportLab's Paragraph supports a ``<link href="...">label</link>`` tag
    that produces a real PDF hyperlink annotation - the user can ctrl-
    click it inside any modern PDF viewer and jump to the URL. We escape
    HTML first, then layer the inline markdown rules + the link
    transformation on top so URLs round-trip correctly.
    """
    escaped = _escape_html(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"\*(.+?)\*", r"<i>\1</i>", escaped)
    escaped = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", escaped)
    # ``[label](url)`` markdown link -> reportlab <link>
    escaped = re.sub(
        r"\[([^\]]+)\]\((https?://[^\s)]+)\)",
        rf'<link href="\2" color="{_LINK_COLOR}"><u>\1</u></link>',
        escaped,
    )
    # Bare URL fallback. Skip URLs already inside a <link href="..."> tag.
    escaped = re.sub(
        r'(?<!href=")(?<!href=\")(https?://[^\s<]+)',
        rf'<link href="\1" color="{_LINK_COLOR}"><u>\1</u></link>',
        escaped,
    )
    return escaped


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _blocks_to_html(blocks: Iterable[Block]) -> str:
    out: list[str] = []
    for block in blocks:
        if block.kind in ("h1", "h2", "h3"):
            level = block.kind[1]
            out.append(f"<h{level}>{_escape_html(block.text)}</h{level}>")
        elif block.kind == "p":
            out.append(f"<p>{_escape_html(block.text)}</p>")
        elif block.kind == "ul":
            inner = "".join(f"<li>{_escape_html(item)}</li>" for item in (block.items or []))
            out.append(f"<ul>{inner}</ul>")
        elif block.kind == "ol":
            inner = "".join(f"<li>{_escape_html(item)}</li>" for item in (block.items or []))
            out.append(f"<ol>{inner}</ol>")
    return "\n".join(out)


# ---------------------------------------------------------------------------
# Summary HTML - the rich one-pager produced when the user clicks
# "Save complete analysis". Pulls the cached structured outputs the
# pipeline already paid for so we do not run another LLM call.
# ---------------------------------------------------------------------------
_SUMMARY_LABELS: dict[str, dict[str, str]] = {
    "en": {
        "page_title": "Application summary",
        "headline": "Application summary",
        "for_role": "{role} at {company}",
        "for_role_no_company": "{role}",
        "match_score_badge": "Match score: {score} / 100",
        "match_report_h": "Match report",
        "verdict": "Verdict",
        "overall_score": "Overall fit",
        "categories": "Score by category",
        "matches": "Matched requirements",
        "gaps": "Risky gaps",
        "ats_present": "ATS keywords present",
        "ats_missing": "ATS keywords missing",
        "interview_questions": "Interview preparation",
        "why_asked": "Why this is asked",
        "suggested_answer": "Suggested answer",
        "skill_gap_plan": "Skill gap plan",
        "criticality": "Importance",
        "why_it_matters": "Why it matters",
        "learning_path": "Learning path",
        "suggested_project": "Suggested project",
        "weeks": "weeks",
        "criticality_critical": "Critical",
        "criticality_important": "Important",
        "criticality_nice_to_have": "Nice to have",
        "tailored_cv_h": "Tailored CV",
        "cover_letter_h": "Cover letter",
        "match_summary_h": "Summary",
        "no_data": "No analysis data captured for this section.",
        "generated_on": "Generated on {timestamp}",
    },
    "cs": {
        "page_title": "Souhrn přihlášky",
        "headline": "Souhrn přihlášky",
        "for_role": "{role} ve společnosti {company}",
        "for_role_no_company": "{role}",
        "match_score_badge": "Skóre shody: {score} / 100",
        "match_report_h": "Report shody",
        "verdict": "Verdikt",
        "overall_score": "Celková shoda",
        "categories": "Skóre po kategoriích",
        "matches": "Splněné požadavky",
        "gaps": "Rizikové mezery",
        "ats_present": "ATS klíčová slova - přítomná",
        "ats_missing": "ATS klíčová slova - chybějící",
        "interview_questions": "Příprava na pohovor",
        "why_asked": "Proč se na to ptají",
        "suggested_answer": "Návrh odpovědi",
        "skill_gap_plan": "Plán doplnění mezer",
        "criticality": "Důležitost",
        "why_it_matters": "Proč to záleží",
        "learning_path": "Učební cesta",
        "suggested_project": "Návrh projektu",
        "weeks": "týdnů",
        "criticality_critical": "Kritické",
        "criticality_important": "Důležité",
        "criticality_nice_to_have": "Hezké mít",
        "tailored_cv_h": "Životopis na míru",
        "cover_letter_h": "Motivační dopis",
        "match_summary_h": "Shrnutí",
        "no_data": "Pro tuto sekci nejsou data z analýzy.",
        "generated_on": "Vygenerováno {timestamp}",
    },
}


def _summary_labels(lang: str) -> dict[str, str]:
    code = (lang or "en").strip().lower()
    return _SUMMARY_LABELS.get(code) or _SUMMARY_LABELS["en"]


def _summary_inline_html(text: str) -> str:
    """Escape ``text`` then keep markdown-bold (``**...**``) as ``<strong>``."""
    escaped = _escape_html(text or "")
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)


def _summary_css(theme: object) -> str:
    """Build a self-contained CSS body that picks up the active palette."""
    accent = getattr(theme, "accent", "#0E7490")
    accent_dark = getattr(theme, "accent_dark", "#0F766E")
    accent_soft = getattr(theme, "accent_soft", "#7DD3FC")
    rule = getattr(theme, "rule", "#14B8A6")
    text_primary = getattr(theme, "text_primary", "#0F172A")
    text_muted = getattr(theme, "text_muted", "#64748B")
    body_font = getattr(
        theme,
        "body_font",
        "'Inter','Segoe UI','Helvetica Neue',Arial,sans-serif",
    )
    heading_font = getattr(theme, "heading_font", body_font)

    return f"""
*{{box-sizing:border-box;margin:0;padding:0;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
::selection{{background:#FDE68A;color:#0F172A}}
@media print{{a[href]::after{{content:""}}}}
html,body{{font-family:{body_font};color:{text_primary};background:#F8FAFC;line-height:1.55;font-size:14px}}
.shell{{max-width:920px;margin:0 auto;padding:32px 28px 64px 28px}}
.banner{{background:linear-gradient(135deg,{accent} 0%,{accent_dark} 100%);color:#FFFFFF;border-radius:14px;padding:28px 30px;margin-bottom:24px;box-shadow:0 18px 40px rgba(15,23,42,0.10)}}
.banner h1{{font-family:{heading_font};font-size:26px;font-weight:800;letter-spacing:-0.005em;margin-bottom:6px}}
.banner .meta{{color:{accent_soft};font-size:14px}}
.banner .stamp{{color:rgba(255,255,255,0.65);font-size:12px;margin-top:10px}}
.scoreboard{{display:flex;flex-wrap:wrap;gap:16px;margin:18px 0 4px 0}}
.score-card{{background:rgba(255,255,255,0.10);border:1px solid rgba(255,255,255,0.20);border-radius:10px;padding:14px 18px;color:#FFFFFF;flex:0 0 auto;min-width:140px}}
.score-card .label{{font-size:11px;text-transform:uppercase;letter-spacing:0.12em;color:{accent_soft};margin-bottom:4px}}
.score-card .value{{font-size:22px;font-weight:800}}
.score-card.verdict{{flex:1;min-width:260px}}
.score-card.verdict .value{{font-size:14px;font-weight:600;line-height:1.45}}
.toc{{background:#FFFFFF;border:1px solid #E2E8F0;border-radius:12px;padding:14px 22px;margin-bottom:18px;box-shadow:0 4px 14px rgba(15,23,42,0.04)}}
.toc ol{{margin:0;padding-left:18px;display:flex;flex-wrap:wrap;gap:4px 22px}}
.toc li{{color:{accent_dark};font-weight:600;font-size:13px}}
.toc a{{color:inherit;text-decoration:none;border-bottom:1px dotted {accent_soft}}}
.toc a:hover{{color:{accent}}}
section.block{{background:#FFFFFF;border:1px solid #E2E8F0;border-radius:12px;padding:22px 24px;margin-bottom:18px;box-shadow:0 4px 14px rgba(15,23,42,0.04)}}
section.block h2{{font-family:{heading_font};font-size:16px;font-weight:800;color:{accent_dark};letter-spacing:0.04em;text-transform:uppercase;border-bottom:2px solid {rule};padding-bottom:6px;margin-bottom:14px}}
section.block h3{{font-family:{heading_font};font-size:14.5px;color:{accent_dark};font-weight:800;letter-spacing:0.02em;margin-top:14px;margin-bottom:6px}}
section.block h4{{font-family:{heading_font};font-size:13.5px;color:{text_primary};font-weight:800;margin-top:8px;margin-bottom:4px}}
.markdown-body{{font-size:14px;line-height:1.6;color:{text_primary}}}
.markdown-body p{{margin:0 0 10px 0}}
.markdown-body h1{{font-family:{heading_font};font-size:20px;color:{accent_dark};font-weight:800;margin:6px 0 10px 0}}
.markdown-body h2{{font-family:{heading_font};font-size:15.5px;color:{accent_dark};font-weight:800;letter-spacing:0.02em;text-transform:uppercase;border-bottom:1px solid {accent_soft};padding-bottom:3px;margin:18px 0 8px 0}}
.markdown-body h3{{font-family:{heading_font};font-size:14px;color:{text_primary};font-weight:800;margin:14px 0 4px 0}}
.markdown-body strong{{color:{accent_dark}}}
.markdown-body ul,.markdown-body ol{{padding-left:22px;margin-bottom:10px}}
.markdown-body li{{margin-bottom:3px}}
.markdown-body em{{color:{text_muted};font-style:italic}}
.markdown-body code{{background:#F1F5F9;padding:1px 5px;border-radius:4px;font-size:0.92em}}
.markdown-body a{{color:{accent_dark};text-decoration:none;border-bottom:1px dotted {accent_soft}}}
.cat-row{{display:flex;align-items:center;gap:14px;margin-bottom:10px}}
.cat-row .name{{flex:0 0 30%;font-weight:600;color:{text_primary}}}
.cat-row .bar{{flex:1;background:#F1F5F9;border-radius:8px;height:8px;overflow:hidden}}
.cat-row .bar .fill{{height:100%;background:linear-gradient(90deg,{accent} 0%,{accent_dark} 100%)}}
.cat-row .num{{flex:0 0 50px;text-align:right;font-variant-numeric:tabular-nums;color:{text_muted};font-weight:600}}
.list-grid{{display:grid;grid-template-columns:1fr 1fr;gap:6px 24px}}
.list-grid li{{margin-left:18px}}
ul.compact li{{margin-left:18px;margin-bottom:4px}}
.chip-row{{display:flex;flex-wrap:wrap;gap:6px}}
.chip{{font-size:12px;padding:3px 10px;border-radius:14px;border:1px solid {accent_soft};color:{accent_dark};background:#F0F9FF;font-weight:600}}
.chip.miss{{border-color:#FBCFE8;color:#9D174D;background:#FDF2F8}}
.q-card{{border-left:3px solid {rule};padding:12px 14px;margin-bottom:14px;background:#F8FAFC;border-radius:0 8px 8px 0}}
.q-card .q{{font-weight:700;color:{text_primary};font-size:15px;margin-bottom:6px}}
.q-card .label{{font-size:11px;text-transform:uppercase;letter-spacing:0.10em;color:{accent_dark};font-weight:700;margin-top:8px;margin-bottom:2px}}
.q-card .body{{color:{text_primary};font-size:13.5px;line-height:1.5}}
.gap-card{{border:1px solid #E2E8F0;border-radius:10px;padding:14px 16px;margin-bottom:14px;background:#FFFFFF}}
.gap-card .head{{display:flex;align-items:baseline;justify-content:space-between;gap:8px;margin-bottom:6px;flex-wrap:wrap}}
.gap-card .name{{font-weight:700;color:{text_primary};font-size:15px}}
.gap-card .meta{{font-size:11.5px;color:{text_muted};display:flex;gap:8px;flex-wrap:wrap;align-items:center}}
.crit-badge{{font-size:10.5px;text-transform:uppercase;letter-spacing:0.08em;font-weight:800;padding:2px 8px;border-radius:10px;border:1px solid transparent}}
.crit-critical{{background:#FEF2F2;border-color:#FECACA;color:#991B1B}}
.crit-important{{background:#FFFBEB;border-color:#FDE68A;color:#92400E}}
.crit-nice_to_have{{background:#F0FDF4;border-color:#BBF7D0;color:#166534}}
.gap-card .label{{font-size:11px;text-transform:uppercase;letter-spacing:0.10em;color:{accent_dark};font-weight:700;margin-top:8px;margin-bottom:2px}}
.gap-card .body{{color:{text_primary};font-size:13.5px;line-height:1.5}}
.gap-card ol{{padding-left:22px;color:{text_primary}}}
.gap-card ol li{{margin-bottom:3px}}
.empty{{color:{text_muted};font-style:italic}}
""".strip()


def _summary_categories_html(categories: list, labels: dict[str, str]) -> str:
    if not categories:
        return f'<div class="empty">{_escape_html(labels["no_data"])}</div>'
    rows: list[str] = []
    for cat in categories:
        if not isinstance(cat, dict):
            continue
        name = _escape_html(str(cat.get("name") or ""))
        score = int(cat.get("score") or 0)
        score = max(0, min(100, score))
        rows.append(
            '<div class="cat-row">'
            f'<div class="name">{name}</div>'
            f'<div class="bar"><div class="fill" style="width:{score}%"></div></div>'
            f'<div class="num">{score}</div>'
            "</div>"
        )
    return "\n".join(rows)


def _summary_list_grid(items: list, *, empty_label: str) -> str:
    items = [str(it).strip() for it in (items or []) if str(it).strip()]
    if not items:
        return f'<div class="empty">{_escape_html(empty_label)}</div>'
    return (
        '<ul class="compact list-grid">'
        + "".join(f"<li>{_summary_inline_html(it)}</li>" for it in items)
        + "</ul>"
    )


def _summary_chip_row(items: list, *, miss: bool, empty_label: str) -> str:
    items = [str(it).strip() for it in (items or []) if str(it).strip()]
    if not items:
        return f'<div class="empty">{_escape_html(empty_label)}</div>'
    cls = "chip miss" if miss else "chip"
    return (
        '<div class="chip-row">'
        + "".join(f'<span class="{cls}">{_escape_html(it)}</span>' for it in items)
        + "</div>"
    )


def _summary_questions_html(items: list, labels: dict[str, str]) -> str:
    if not items:
        return f'<div class="empty">{_escape_html(labels["no_data"])}</div>'
    cards: list[str] = []
    for idx, raw in enumerate(items, start=1):
        # Tolerate the legacy ``list[str]`` shape from older saved runs:
        # if the item is a bare string we render it as the question with
        # empty rationale / answer fields rather than crashing.
        if isinstance(raw, str):
            question = raw.strip()
            why = ""
            answer = ""
        elif isinstance(raw, dict):
            question = str(raw.get("question") or "").strip()
            why = str(raw.get("why_asked") or "").strip()
            answer = str(raw.get("suggested_answer") or "").strip()
        else:
            continue
        if not question:
            continue
        body_parts: list[str] = [
            f'<div class="q">{idx}. {_summary_inline_html(question)}</div>'
        ]
        if why:
            body_parts.append(
                f'<div class="label">{_escape_html(labels["why_asked"])}</div>'
                f'<div class="body">{_summary_inline_html(why)}</div>'
            )
        if answer:
            body_parts.append(
                f'<div class="label">{_escape_html(labels["suggested_answer"])}</div>'
                f'<div class="body">{_summary_inline_html(answer)}</div>'
            )
        cards.append(f'<div class="q-card">{"".join(body_parts)}</div>')
    return "\n".join(cards)


def _summary_skill_gaps_html(items: list, labels: dict[str, str]) -> str:
    if not items:
        return f'<div class="empty">{_escape_html(labels["no_data"])}</div>'
    crit_order = {"critical": 0, "important": 1, "nice_to_have": 2}
    sorted_items = sorted(
        (it for it in items if isinstance(it, dict)),
        key=lambda it: crit_order.get(
            str(it.get("criticality") or "").lower(), 99
        ),
    )
    cards: list[str] = []
    for entry in sorted_items:
        skill = str(entry.get("skill") or "").strip()
        if not skill:
            continue
        action = str(entry.get("action") or "").strip()
        weeks = entry.get("timeline_weeks")
        criticality = str(entry.get("criticality") or "").strip().lower()
        why = str(entry.get("why_it_matters") or "").strip()
        suggested_project = str(entry.get("suggested_project") or "").strip()
        learning_path_raw = entry.get("learning_path") or []
        learning_path = [
            str(step).strip() for step in learning_path_raw if str(step).strip()
        ]

        crit_label = ""
        if criticality:
            label_key = f"criticality_{criticality}"
            crit_label = labels.get(label_key) or criticality.replace("_", " ").title()

        meta_bits: list[str] = []
        if crit_label:
            crit_class = (
                "crit-badge"
                f" crit-{criticality}"
                if criticality in {"critical", "important", "nice_to_have"}
                else "crit-badge"
            )
            meta_bits.append(
                f'<span class="{crit_class}">{_escape_html(crit_label)}</span>'
            )
        if isinstance(weeks, int) and weeks > 0:
            meta_bits.append(
                f'<span>~{weeks} {_escape_html(labels["weeks"])}</span>'
            )

        head = (
            '<div class="head">'
            f'<div class="name">{_summary_inline_html(skill)}</div>'
            f'<div class="meta">{"".join(meta_bits)}</div>'
            "</div>"
        )

        body_parts: list[str] = [head]
        if action:
            body_parts.append(f'<div class="body">{_summary_inline_html(action)}</div>')
        if why:
            body_parts.append(
                f'<div class="label">{_escape_html(labels["why_it_matters"])}</div>'
                f'<div class="body">{_summary_inline_html(why)}</div>'
            )
        if learning_path:
            body_parts.append(
                f'<div class="label">{_escape_html(labels["learning_path"])}</div>'
                "<ol>"
                + "".join(
                    f"<li>{_summary_inline_html(step)}</li>"
                    for step in learning_path
                )
                + "</ol>"
            )
        if suggested_project:
            body_parts.append(
                f'<div class="label">{_escape_html(labels["suggested_project"])}</div>'
                f'<div class="body">{_summary_inline_html(suggested_project)}</div>'
            )
        cards.append(f'<div class="gap-card">{"".join(body_parts)}</div>')
    return "\n".join(cards)


def _summary_markdown_to_html(text: str) -> str:
    """Convert markdown into a styled, escaped HTML fragment.

    Reuses the existing :func:`parse_markdown` block parser so headings,
    bullet lists, ordered lists, and paragraphs are all preserved; bold
    + italics + inline code are converted via the same regex used by
    the markdown HTML exporter.
    """
    if not text or not text.strip():
        return ""
    blocks = parse_markdown(text)
    parts: list[str] = []
    for block in blocks:
        if block.kind in ("h1", "h2", "h3"):
            level = block.kind[1]
            parts.append(
                f"<h{level}>{_summary_inline_md_html(block.text)}</h{level}>"
            )
        elif block.kind == "p":
            parts.append(f"<p>{_summary_inline_md_html(block.text)}</p>")
        elif block.kind == "ul":
            inner = "".join(
                f"<li>{_summary_inline_md_html(item)}</li>"
                for item in (block.items or [])
            )
            parts.append(f"<ul>{inner}</ul>")
        elif block.kind == "ol":
            inner = "".join(
                f"<li>{_summary_inline_md_html(item)}</li>"
                for item in (block.items or [])
            )
            parts.append(f"<ol>{inner}</ol>")
    if not parts:
        return f"<p>{_summary_inline_md_html(text)}</p>"
    return f'<div class="markdown-body">{chr(10).join(parts)}</div>'


def _summary_inline_md_html(text: str) -> str:
    """Inline markdown rules used inside the summary's markdown blocks.

    Stricter than :func:`_summary_inline_html` because the saved
    documents may contain ``*italic*`` and inline ``code`` we want to
    preserve, plus naked URLs that should turn into clickable links.
    """
    escaped = _escape_html(text or "")
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", escaped)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    # Markdown links: [label](url)
    escaped = re.sub(
        r"\[([^\]]+)\]\((https?://[^\s)]+)\)",
        r'<a href="\2">\1</a>',
        escaped,
    )
    # Naked URLs left over after the markdown link pass.
    escaped = re.sub(
        r"(?<!href=\")(https?://[^\s<]+)",
        r'<a href="\1">\1</a>',
        escaped,
    )
    return escaped


def build_summary_html(
    *,
    candidate: dict | None,
    job_spec: dict | None,
    match: dict | None,
    theme: object,
    output_lang: str = "en",
    documents: dict[str, str] | None = None,
    timestamp: str = "",
) -> str:
    """Assemble a rich, themed summary page from cached structured outputs.

    The output is a single self-contained HTML document with inline
    CSS, suitable for "Open in browser" or for archival inside the
    saved run folder. No LLM call - this is purely a deterministic
    render of the JSON payloads + saved markdown documents we already
    paid for.

    ``documents`` (optional) is a copy of ``STATE.documents`` mapping
    doc-kind -> markdown body. The Tailored CV and Cover Letter
    bodies, when present, are inlined into their own sections so the
    user has the entire application packet in one HTML file - matching
    the reference Application_Summary.html layout.

    ``timestamp`` is rendered into the banner stamp; pass an empty
    string to skip the stamp.
    """
    candidate = candidate or {}
    job_spec = job_spec or {}
    match = match or {}
    documents = documents or {}
    labels = _summary_labels(output_lang)
    css = _summary_css(theme)
    lang_attr = (output_lang or "en").strip().lower() or "en"

    role = str(job_spec.get("title") or "").strip()
    company = str(job_spec.get("company") or "").strip()
    if role and company:
        sub = labels["for_role"].format(
            role=_escape_html(role), company=_escape_html(company)
        )
    elif role:
        sub = labels["for_role_no_company"].format(role=_escape_html(role))
    else:
        sub = ""

    overall = int(match.get("overall_score") or 0)
    overall = max(0, min(100, overall))
    verdict = str(match.get("verdict") or "").strip()

    scoreboard = (
        '<div class="scoreboard">'
        '<div class="score-card">'
        f'<div class="label">{_escape_html(labels["overall_score"])}</div>'
        f'<div class="value">{overall}/100</div>'
        "</div>"
        + (
            '<div class="score-card verdict">'
            f'<div class="label">{_escape_html(labels["verdict"])}</div>'
            f'<div class="value">{_summary_inline_html(verdict)}</div>'
            "</div>"
            if verdict
            else ""
        )
        + "</div>"
    )
    stamp = (
        f'<div class="stamp">{_escape_html(labels["generated_on"].format(timestamp=timestamp))}</div>'
        if timestamp
        else ""
    )

    # Each section gets its own anchor so the table-of-contents links
    # work and the user can jump between sections in a long page.
    sections: list[tuple[str, str, str]] = []  # (anchor, label, html)

    sections.append(
        (
            "match-report",
            labels["match_report_h"],
            (
                f'<h3>{_escape_html(labels["categories"])}</h3>'
                + _summary_categories_html(match.get("categories") or [], labels)
                + f'<h3>{_escape_html(labels["matches"])}</h3>'
                + _summary_list_grid(
                    match.get("matches") or [], empty_label=labels["no_data"]
                )
                + f'<h3>{_escape_html(labels["gaps"])}</h3>'
                + _summary_list_grid(
                    match.get("gaps") or [], empty_label=labels["no_data"]
                )
                + f'<h3>{_escape_html(labels["ats_present"])}</h3>'
                + _summary_chip_row(
                    match.get("ats_keywords_present") or [],
                    miss=False,
                    empty_label=labels["no_data"],
                )
                + f'<h3>{_escape_html(labels["ats_missing"])}</h3>'
                + _summary_chip_row(
                    match.get("ats_keywords_missing") or [],
                    miss=True,
                    empty_label=labels["no_data"],
                )
                + (
                    f'<h3>{_escape_html(labels["match_summary_h"])}</h3>'
                    f'<p class="markdown-body">{_summary_inline_html(verdict)}</p>'
                    if verdict
                    else ""
                )
            ),
        )
    )

    tailored_cv = (documents or {}).get("tailored_cv", "")
    if tailored_cv and tailored_cv.strip():
        sections.append(
            (
                "tailored-cv",
                labels["tailored_cv_h"],
                _summary_markdown_to_html(tailored_cv),
            )
        )

    cover_letter = (documents or {}).get("cover_letter", "")
    if cover_letter and cover_letter.strip():
        sections.append(
            (
                "cover-letter",
                labels["cover_letter_h"],
                _summary_markdown_to_html(cover_letter),
            )
        )

    sections.append(
        (
            "interview",
            labels["interview_questions"],
            _summary_questions_html(
                match.get("interview_questions") or [], labels
            ),
        )
    )
    sections.append(
        (
            "skill-gap",
            labels["skill_gap_plan"],
            _summary_skill_gaps_html(
                match.get("skill_gap_plan") or [], labels
            ),
        )
    )

    toc_html = (
        '<div class="toc"><ol>'
        + "".join(
            f'<li><a href="#{anchor}">{_escape_html(label)}</a></li>'
            for anchor, label, _ in sections
        )
        + "</ol></div>"
    )

    sections_html = "\n".join(
        f'<section class="block" id="{anchor}">'
        f'<h2>{_escape_html(label)}</h2>'
        f"{body}"
        "</section>"
        for anchor, label, body in sections
    )

    return (
        "<!doctype html>\n"
        f'<html lang="{lang_attr}"><head><meta charset="utf-8"/>'
        f"<title>{_escape_html(labels['page_title'])}</title>"
        f"<style>{css}</style></head>"
        '<body><div class="shell">'
        '<div class="banner">'
        f'<h1>{_escape_html(labels["headline"])}</h1>'
        + (f'<div class="meta">{sub}</div>' if sub else "")
        + scoreboard
        + stamp
        + "</div>"
        + toc_html
        + sections_html
        + "</div></body></html>"
    )
