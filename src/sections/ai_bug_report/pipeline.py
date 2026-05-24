"""Orchestration for the AI Bug Report section.

Two public entry points:

* :func:`generate_bug_report` - one structured AI call that turns the
  description + environment hints + supporting docs + screenshots into
  a strict ``BUG_REPORT_SCHEMA`` JSON payload. Screenshots are sent
  through :func:`src.services.ai_provider.run`'s ``images=`` kwarg so
  the model can read them via its native vision API.

* :func:`save_bug_report_docx` - writes the saved-run folder under
  ``outputs/ai_bug_report/<title-slug>-<timestamp>/`` with:
    * ``<Title>.docx`` - heading, env table, numbered STR, expected vs
      actual, embedded screenshots, attachments summary,
    * ``<Title>.md`` - plain-text mirror,
    * ``summary.json`` - the schema dict + run metadata,
    * a ``RunSummary(note="ai_bug_report")`` entry in
      ``~/AI Hub/history.json``.

Both functions short-circuit in demo mode and emit structured
``*_start`` / ``*_done`` / ``*_failed`` log lines per
``ai-section.mdc`` so the user can debug from Settings -> Debug logs.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from src.services import ai_provider, store
from src.services import logger as logger_service
from src.services.cost_tracker import COST
from src.sections.ai_bug_report import prompts, schema
from src.sections.ai_bug_report.refs import REFS
from src.sections.ai_bug_report.state import (
    SEVERITY_VALUES,
    PRIORITY_VALUES,
    REPRODUCIBILITY_VALUES,
    STATE,
)


# --- Result type -------------------------------------------------------------


@dataclass
class StepResult:
    ok: bool
    error: str = ""


@dataclass
class SaveResult:
    ok: bool
    folder: str = ""
    path: str = ""
    error: str = ""


# --- Demo data ---------------------------------------------------------------


_DEMO_REPORT: dict[str, Any] = {
    "title": "Profile settings fail to save and avatar upload breaks",
    "summary": (
        "Two distinct profile settings regressions are visible in the sample "
        "inputs: saving the display name never completes, and uploading a new "
        "avatar fails validation after the image is selected."
    ),
    "scenarios": [
        {
            "title": "Profile Save button stays in loading state on Chrome 124",
            "summary": (
                "Saving the user profile on the Settings page leaves the Save "
                "button stuck in the spinner state. The change is not persisted "
                "and no toast is shown to the user."
            ),
            "severity": "High",
            "priority": "P1",
            "reproducibility": "Always",
            "environment": {
                "browser": "Chrome 124.0.6367.92",
                "os": "Windows 11 23H2",
                "device": "Desktop",
                "app_version": "2026.05.4",
                "url": "https://app.example.com/settings/profile",
            },
            "preconditions": [
                "User is logged in with a Standard plan account.",
                "Profile already has a display name and an avatar.",
            ],
            "steps_to_reproduce": [
                "Open https://app.example.com/settings/profile in Chrome.",
                "Change the Display name field to any non-empty value.",
                "Click the Save button.",
                "Wait 10 seconds.",
            ],
            "expected_result": (
                "The Save button returns to its idle state, a green 'Profile "
                "saved' toast appears, and the new display name is shown in the "
                "top-right user menu."
            ),
            "actual_result": (
                "The Save button stays in its loading (spinner) state indefinitely. "
                "No toast is shown. Reloading the page reverts the display name "
                "to the previous value, so the change was not persisted."
            ),
            "additional_notes": (
                "(inferred) The 500 in the console suggests a server-side "
                "regression in PUT /api/profile rather than a frontend bug."
            ),
        },
        {
            "title": "Avatar upload shows validation error for valid PNG",
            "summary": (
                "A valid PNG avatar appears to upload, but the form shows a "
                "validation error and the image is not applied to the profile."
            ),
            "severity": "Medium",
            "priority": "P2",
            "reproducibility": "Sometimes",
            "environment": {
                "browser": "Chrome 124.0.6367.92",
                "os": "Windows 11 23H2",
                "device": "Desktop",
                "app_version": "2026.05.4",
                "url": "https://app.example.com/settings/profile",
            },
            "preconditions": [
                "User is logged in with a Standard plan account.",
                "User has a PNG avatar under the documented 2 MB limit.",
            ],
            "steps_to_reproduce": [
                "Open https://app.example.com/settings/profile in Chrome.",
                "Click Change avatar.",
                "Select a valid PNG image under 2 MB.",
                "Wait for the upload validation to finish.",
            ],
            "expected_result": "The avatar preview updates and the user can save the profile.",
            "actual_result": (
                "The form shows a validation error even though the file meets "
                "the documented PNG and size requirements."
            ),
            "additional_notes": (
                "(inferred) This may share the same profile endpoint regression "
                "as the display-name save failure."
            ),
        },
    ],
    "attachments_summary": [
        "Screenshot 1: profile page with the Save button in the spinner state.",
        "Document 1: relevant browser console output showing a 500 from PUT /api/profile.",
    ],
    "additional_notes": (
        "(inferred) Both scenarios touch profile persistence. Worth attaching "
        "the request payloads from DevTools when reproducing."
    ),
}


def load_demo() -> None:
    STATE.last_report = dict(_DEMO_REPORT)
    STATE.last_error = ""
    STATE.activity = "ready"
    REFS.request_context_refresh()


# --- Helpers -----------------------------------------------------------------


def _set_activity(value: str) -> None:
    """Update the right-hand activity badge from any thread."""
    prev = STATE.activity
    STATE.activity = value
    if prev != value:
        logger_service.log_event(
            "INFO",
            "ai_bug_report.pipeline",
            "activity_change",
            prev=prev,
            new=value,
        )
    REFS.request_context_refresh()


def _normalise_enum(value: Any, allowed: tuple, default: str) -> str:
    candidate = str(value or "").strip()
    for item in allowed:
        if candidate.lower() == item.lower():
            return item
    return default


def _normalise_scenario(data: Any, *, fallback_title: str = "Untitled bug") -> dict[str, Any]:
    """Coerce one scenario into the shape the UI / DOCX expects."""
    if not isinstance(data, dict):
        return {}

    out: dict[str, Any] = {}
    out["title"] = str(data.get("title") or "").strip() or fallback_title
    out["summary"] = str(data.get("summary") or "").strip()
    out["severity"] = _normalise_enum(
        data.get("severity"), SEVERITY_VALUES, "Medium"
    )
    out["priority"] = _normalise_enum(
        data.get("priority"), PRIORITY_VALUES, "P2"
    )
    out["reproducibility"] = _normalise_enum(
        data.get("reproducibility"), REPRODUCIBILITY_VALUES, "Unknown"
    )

    env = data.get("environment") or {}
    if not isinstance(env, dict):
        env = {}
    out["environment"] = {
        "browser": str(env.get("browser") or "").strip(),
        "os": str(env.get("os") or "").strip(),
        "device": str(env.get("device") or "").strip(),
        "app_version": str(env.get("app_version") or "").strip(),
        "url": str(env.get("url") or "").strip(),
    }

    out["preconditions"] = [
        str(item).strip()
        for item in (data.get("preconditions") or [])
        if str(item).strip()
    ]
    out["steps_to_reproduce"] = [
        str(item).strip()
        for item in (data.get("steps_to_reproduce") or [])
        if str(item).strip()
    ]
    out["expected_result"] = str(data.get("expected_result") or "").strip()
    out["actual_result"] = str(data.get("actual_result") or "").strip()
    out["additional_notes"] = str(data.get("additional_notes") or "").strip()
    return out


def _normalise_report(data: Any) -> dict[str, Any]:
    """Normalise a multi-scenario report, accepting legacy single reports."""
    if not isinstance(data, dict):
        return {}

    raw_scenarios = data.get("scenarios")
    if isinstance(raw_scenarios, list) and raw_scenarios:
        scenarios = [
            scenario
            for scenario in (
                _normalise_scenario(item, fallback_title=f"Scenario {idx}")
                for idx, item in enumerate(raw_scenarios, start=1)
            )
            if scenario
        ]
    else:
        # Backwards compatibility for reports saved before ``scenarios[]``
        # existed. Treat the old flat object as a single scenario.
        legacy = _normalise_scenario(
            data,
            fallback_title=str(data.get("title") or "").strip() or "Untitled bug",
        )
        scenarios = [legacy] if legacy else []

    if not scenarios:
        scenarios = [
            _normalise_scenario(
                {"title": "Untitled bug", "steps_to_reproduce": []},
                fallback_title="Untitled bug",
            )
        ]

    title = str(data.get("title") or "").strip() or scenarios[0].get("title") or "Bug report"
    summary = str(data.get("summary") or "").strip() or scenarios[0].get("summary", "")
    attachments_summary = [
        str(item).strip()
        for item in (data.get("attachments_summary") or [])
        if str(item).strip()
    ]

    return {
        "title": title,
        "summary": summary,
        "scenarios": scenarios,
        "attachments_summary": attachments_summary,
        "additional_notes": str(data.get("additional_notes") or "").strip(),
    }


def _primary_scenario(report: dict[str, Any]) -> dict[str, Any]:
    scenarios = report.get("scenarios") or []
    if isinstance(scenarios, list) and scenarios:
        first = scenarios[0]
        if isinstance(first, dict):
            return first
    return _normalise_scenario(report)


# --- AI call -----------------------------------------------------------------


def generate_followup_questions(*, output_lang: str) -> StepResult:
    """Ask the LLM what facts it is missing before filing the report.

    Returns ``StepResult(ok=True)`` on success and stores the cleaned
    list in ``STATE.followup_questions`` (may be an empty list when
    the inputs are already good enough). Demo mode and zero-input
    short-circuit so the call never wastes tokens.
    """
    logger_service.log_event(
        "INFO",
        "ai_bug_report.pipeline",
        "generate_followup_questions_start",
        output_lang=output_lang,
        demo_mode=STATE.demo_mode,
        description_chars=len(STATE.description or ""),
        images=len(STATE.images),
        documents=len(STATE.documents),
    )

    if STATE.demo_mode:
        STATE.followup_questions = []
        logger_service.log_event(
            "INFO",
            "ai_bug_report.pipeline",
            "generate_followup_questions_demo_done",
        )
        return StepResult(ok=True)

    if not STATE.can_generate():
        STATE.followup_questions = []
        return StepResult(ok=True)

    _set_activity("followups")

    documents_payload = [
        {"name": d.name, "ext": d.ext, "text": d.text}
        for d in STATE.documents
        if d.text
    ]
    image_payload = [
        {"data": img.bytes_data, "mime": img.mime}
        for img in STATE.images
        if img.bytes_data
    ]
    user_prompt = prompts.build_followup_user(
        description=STATE.description,
        environment_hint=STATE.environment_hint,
        documents=documents_payload,
        image_count=len(image_payload),
        output_lang=output_lang,
    )

    try:
        result = ai_provider.run(
            system=prompts.FOLLOWUP_QUESTIONS_SYSTEM,
            user=user_prompt,
            schema=schema.FOLLOWUP_QUESTIONS_SCHEMA,
            schema_name="bug_report_followup_questions",
            max_output_tokens=1200,
            temperature=0.2,
            images=image_payload or None,
        )
    except ai_provider.ProviderError as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline",
            "generate_followup_questions_provider_error",
            exc,
        )
        STATE.last_error = str(exc)
        _set_activity("error")
        return StepResult(ok=False, error=str(exc))
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "generate_followup_questions_failed", exc,
        )
        STATE.last_error = f"Follow-up questions failed: {exc}"
        _set_activity("error")
        return StepResult(ok=False, error=STATE.last_error)

    if not isinstance(result.data, dict):
        STATE.last_error = "AI did not return a follow-up questions JSON."
        _set_activity("error")
        return StepResult(ok=False, error=STATE.last_error)

    raw = result.data.get("questions") or []
    cleaned: list[dict] = []
    for q in raw:
        if not isinstance(q, dict):
            continue
        topic = str(q.get("topic") or "").strip()
        question = str(q.get("question") or "").strip()
        if not topic or not question:
            continue
        options = [
            str(opt).strip()
            for opt in (q.get("options") or [])
            if isinstance(opt, (str, int, float)) and str(opt).strip()
        ]
        cleaned.append(
            {
                "topic": topic,
                "question": question,
                "rationale": str(q.get("rationale") or "").strip(),
                "options": options,
                "multi_select": bool(q.get("multi_select")),
                "allow_free_text": bool(q.get("allow_free_text", True)),
            }
        )
    STATE.followup_questions = cleaned
    logger_service.log_event(
        "INFO",
        "ai_bug_report.pipeline",
        "generate_followup_questions_done",
        count=len(cleaned),
    )
    return StepResult(ok=True)


def generate_bug_report(*, output_lang: str) -> StepResult:
    """Run the structured-output AI call with the current STATE.

    Returns ``StepResult(ok=True)`` on success and stores the parsed
    payload in ``STATE.last_report``. Demo mode short-circuits and
    loads :data:`_DEMO_REPORT` without spending tokens. The optional
    ``STATE.followup_qa`` answers (collected from the clarifying-
    questions modal) are forwarded to the prompt builder as ground
    truth.
    """
    logger_service.log_event(
        "INFO",
        "ai_bug_report.pipeline",
        "generate_bug_report_start",
        output_lang=output_lang,
        demo_mode=STATE.demo_mode,
        description_chars=len(STATE.description or ""),
        env_hint_chars=len(STATE.environment_hint or ""),
        images=len(STATE.images),
        documents=len(STATE.documents),
        followups=len(STATE.followup_qa or []),
    )

    if STATE.demo_mode:
        load_demo()
        logger_service.log_event(
            "INFO", "ai_bug_report.pipeline", "generate_bug_report_demo_done"
        )
        return StepResult(ok=True)

    if not STATE.can_generate():
        STATE.last_error = "Add a description or at least one attachment first."
        _set_activity("error")
        return StepResult(ok=False, error=STATE.last_error)

    _set_activity("generating")

    documents_payload = [
        {"name": d.name, "ext": d.ext, "text": d.text}
        for d in STATE.documents
        if d.text
    ]
    image_payload = [
        {"data": img.bytes_data, "mime": img.mime}
        for img in STATE.images
        if img.bytes_data
    ]
    user_prompt = prompts.build_bug_report_user(
        description=STATE.description,
        environment_hint=STATE.environment_hint,
        documents=documents_payload,
        image_count=len(image_payload),
        output_lang=output_lang,
        followup_qa=list(STATE.followup_qa or []),
    )

    try:
        result = ai_provider.run(
            system=prompts.SYSTEM_PROMPT,
            user=user_prompt,
            schema=schema.BUG_REPORT_SCHEMA,
            schema_name="bug_report",
            max_output_tokens=4000,
            temperature=0.2,
            images=image_payload or None,
        )
    except ai_provider.ProviderError as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "generate_bug_report_provider_error", exc,
        )
        STATE.last_error = str(exc)
        _set_activity("error")
        return StepResult(ok=False, error=str(exc))
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "generate_bug_report_failed", exc,
        )
        STATE.last_error = f"Generation failed: {exc}"
        _set_activity("error")
        return StepResult(ok=False, error=STATE.last_error)

    if not isinstance(result.data, dict):
        STATE.last_error = "AI did not return a structured bug report."
        _set_activity("error")
        return StepResult(ok=False, error=STATE.last_error)

    STATE.last_report = _normalise_report(result.data)
    STATE.last_error = ""
    _set_activity("ready")
    logger_service.log_event(
        "INFO",
        "ai_bug_report.pipeline",
        "generate_bug_report_done",
        title_chars=len(STATE.last_report.get("title", "")),
        scenarios=len(STATE.last_report.get("scenarios", [])),
        steps=sum(
            len(item.get("steps_to_reproduce", []))
            for item in STATE.last_report.get("scenarios", [])
            if isinstance(item, dict)
        ),
    )
    return StepResult(ok=True)


# --- DOCX export -------------------------------------------------------------


_SEVERITY_COLOR: dict[str, str] = {
    "Critical": "B91C1C",
    "High": "DC2626",
    "Medium": "D97706",
    "Low": "059669",
}


def _slug(text: str, max_length: int = 40) -> str:
    text = (text or "").strip().lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = text.strip("-")
    if not text:
        text = "bug-report"
    return text[:max_length]


def _safe_filename(text: str, max_length: int = 60) -> str:
    text = (text or "").strip()
    text = re.sub(r"[\\/:*?\"<>|]+", "", text)
    text = re.sub(r"\s+", "_", text)
    return text[:max_length] or "Bug_Report"


def _format_environment(env: dict[str, str], labels: dict[str, str]) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for key, label_key in (
        ("browser", "env_browser"),
        ("os", "env_os"),
        ("device", "env_device"),
        ("app_version", "env_app_version"),
        ("url", "env_url"),
    ):
        value = (env.get(key) or "").strip()
        if value:
            rows.append((labels[label_key], value))
    return rows


def _render_markdown(
    report: dict[str, Any], labels: dict[str, str]
) -> str:
    scenarios = report.get("scenarios")
    if isinstance(scenarios, list) and scenarios:
        parts: list[str] = []
        parts.append(f"# {report.get('title', 'Bug report')}")
        parts.append("")
        if report.get("summary"):
            parts.append(report["summary"])
            parts.append("")
        for idx, scenario in enumerate(scenarios, start=1):
            if not isinstance(scenario, dict):
                continue
            scenario_lines = _render_markdown(scenario, labels).strip().splitlines()
            for line in scenario_lines:
                if line.startswith("# "):
                    parts.append(f"## {labels['preview_scenario_label']} {idx}: {line[2:]}")
                elif line.startswith("## "):
                    parts.append(f"### {line[3:]}")
                else:
                    parts.append(line)
            parts.append("")
        if report.get("attachments_summary"):
            parts.append(f"## {labels['preview_attachments_label']}")
            for item in report["attachments_summary"]:
                parts.append(f"- {item}")
            parts.append("")
        if report.get("additional_notes"):
            parts.append(f"## {labels['preview_notes_label']}")
            parts.append(report["additional_notes"])
            parts.append("")
        return "\n".join(parts).strip() + "\n"

    parts: list[str] = []
    parts.append(f"# {report.get('title', 'Bug report')}")
    parts.append("")
    parts.append(
        f"**{labels['preview_severity_label']}:** {report.get('severity', '')}  "
        f"**{labels['preview_priority_label']}:** {report.get('priority', '')}  "
        f"**{labels['preview_repro_label']}:** {report.get('reproducibility', '')}"
    )
    parts.append("")
    if report.get("summary"):
        parts.append(f"## {labels['preview_summary_label']}")
        parts.append(report["summary"])
        parts.append("")

    env_rows = _format_environment(report.get("environment") or {}, labels)
    if env_rows:
        parts.append(f"## {labels['preview_environment_label']}")
        for name, value in env_rows:
            parts.append(f"- **{name}:** {value}")
        parts.append("")

    if report.get("preconditions"):
        parts.append(f"## {labels['preview_preconditions_label']}")
        for item in report["preconditions"]:
            parts.append(f"- {item}")
        parts.append("")

    if report.get("steps_to_reproduce"):
        parts.append(f"## {labels['preview_str_label']}")
        for idx, step in enumerate(report["steps_to_reproduce"], start=1):
            parts.append(f"{idx}. {step}")
        parts.append("")

    if report.get("expected_result"):
        parts.append(f"## {labels['preview_expected_label']}")
        parts.append(report["expected_result"])
        parts.append("")

    if report.get("actual_result"):
        parts.append(f"## {labels['preview_actual_label']}")
        parts.append(report["actual_result"])
        parts.append("")

    if report.get("attachments_summary"):
        parts.append(f"## {labels['preview_attachments_label']}")
        for item in report["attachments_summary"]:
            parts.append(f"- {item}")
        parts.append("")

    if report.get("additional_notes"):
        parts.append(f"## {labels['preview_notes_label']}")
        parts.append(report["additional_notes"])
        parts.append("")

    return "\n".join(parts).strip() + "\n"


def _write_docx(
    target: Path,
    *,
    report: dict[str, Any],
    labels: dict[str, str],
    images: list,
) -> None:
    """Build the Word document via :mod:`python-docx`.

    Layout:

    * Heading 1 = report title.
    * Severity / Priority / Reproducibility line in bold.
    * Summary paragraph.
    * Environment 2-col table (label, value) - only non-empty rows.
    * Preconditions bullet list.
    * Steps to reproduce numbered list.
    * Expected result heading + paragraph (green-ish accent).
    * Actual result heading + paragraph (red-ish accent).
    * Each attached screenshot embedded inline below a small caption.
    * Attachments summary bullet list.
    * Additional notes paragraph.

    The function imports python-docx lazily so the rest of the
    pipeline keeps working when the library is somehow missing (the
    caller catches the ImportError and surfaces a friendly message).
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
        from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run `pip install -r requirements.txt`."
        ) from exc

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    scenarios = report.get("scenarios")
    if isinstance(scenarios, list) and scenarios:
        document.add_heading(report.get("title") or "Bug report", level=1)
        if report.get("summary"):
            document.add_paragraph(report["summary"])

        for idx, scenario in enumerate(scenarios, start=1):
            if not isinstance(scenario, dict):
                continue
            document.add_heading(
                f"{labels['preview_scenario_label']} {idx}: "
                f"{scenario.get('title') or 'Bug'}",
                level=2,
            )
            meta = document.add_paragraph()
            meta.add_run(f"{labels['preview_severity_label']}: ").bold = True
            sev_run = meta.add_run(scenario.get("severity") or "")
            sev_run.bold = True
            try:
                color_hex = _SEVERITY_COLOR.get(scenario.get("severity", ""), "374151")
                sev_run.font.color.rgb = RGBColor.from_string(color_hex)
            except Exception:
                pass
            meta.add_run(f"   {labels['preview_priority_label']}: ").bold = True
            meta.add_run(scenario.get("priority") or "")
            meta.add_run(f"   {labels['preview_repro_label']}: ").bold = True
            meta.add_run(scenario.get("reproducibility") or "")

            if scenario.get("summary"):
                document.add_heading(labels["preview_summary_label"], level=3)
                document.add_paragraph(scenario["summary"])

            env_rows = _format_environment(scenario.get("environment") or {}, labels)
            if env_rows:
                document.add_heading(labels["preview_environment_label"], level=3)
                table = document.add_table(rows=len(env_rows), cols=2)
                table.style = "Light List Accent 1"
                for row_idx, (name, value) in enumerate(env_rows):
                    cell_label = table.rows[row_idx].cells[0]
                    cell_value = table.rows[row_idx].cells[1]
                    cell_label.text = ""
                    cell_value.text = ""
                    run = cell_label.paragraphs[0].add_run(name)
                    run.bold = True
                    cell_value.paragraphs[0].add_run(value)

            if scenario.get("preconditions"):
                document.add_heading(labels["preview_preconditions_label"], level=3)
                for item in scenario["preconditions"]:
                    document.add_paragraph(item, style="List Bullet")

            if scenario.get("steps_to_reproduce"):
                document.add_heading(labels["preview_str_label"], level=3)
                for step in scenario["steps_to_reproduce"]:
                    document.add_paragraph(step, style="List Number")

            if scenario.get("expected_result"):
                h_para = document.add_heading(labels["preview_expected_label"], level=3)
                try:
                    for run in h_para.runs:
                        run.font.color.rgb = RGBColor.from_string("059669")
                except Exception:
                    pass
                document.add_paragraph(scenario["expected_result"])

            if scenario.get("actual_result"):
                h_para = document.add_heading(labels["preview_actual_label"], level=3)
                try:
                    for run in h_para.runs:
                        run.font.color.rgb = RGBColor.from_string("B91C1C")
                except Exception:
                    pass
                document.add_paragraph(scenario["actual_result"])

            if scenario.get("additional_notes"):
                document.add_heading(labels["preview_notes_label"], level=3)
                document.add_paragraph(scenario["additional_notes"])

        if images:
            document.add_heading(labels["preview_attachments_label"], level=2)
            for idx, img in enumerate(images, start=1):
                caption = document.add_paragraph()
                caption_run = caption.add_run(f"Screenshot {idx}: {img.name}")
                caption_run.bold = True
                try:
                    document.add_picture(img.path, width=Inches(5.5))
                except Exception as exc:
                    logger_service.log_exception(
                        "ai_bug_report.pipeline",
                        "write_docx_add_picture_failed",
                        exc,
                        name=img.name,
                    )
                    document.add_paragraph(f"[Could not embed {img.name}: {exc}]")

        if report.get("attachments_summary"):
            if not images:
                document.add_heading(labels["preview_attachments_label"], level=2)
            for item in report["attachments_summary"]:
                document.add_paragraph(item, style="List Bullet")

        if report.get("additional_notes"):
            document.add_heading(labels["preview_notes_label"], level=2)
            document.add_paragraph(report["additional_notes"])

        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(target))
        return

    title_para = document.add_heading(report.get("title") or "Bug report", level=1)
    title_run = title_para.runs[0] if title_para.runs else None
    if title_run is not None:
        try:
            title_run.font.color.rgb = RGBColor.from_string("1F2937")
        except Exception:
            pass

    meta = document.add_paragraph()
    meta.add_run(f"{labels['preview_severity_label']}: ").bold = True
    sev_run = meta.add_run(report.get("severity") or "")
    sev_run.bold = True
    try:
        color_hex = _SEVERITY_COLOR.get(report.get("severity", ""), "374151")
        sev_run.font.color.rgb = RGBColor.from_string(color_hex)
    except Exception:
        pass
    meta.add_run(f"   {labels['preview_priority_label']}: ").bold = True
    meta.add_run(report.get("priority") or "")
    meta.add_run(f"   {labels['preview_repro_label']}: ").bold = True
    meta.add_run(report.get("reproducibility") or "")

    if report.get("summary"):
        document.add_heading(labels["preview_summary_label"], level=2)
        document.add_paragraph(report["summary"])

    env_rows = _format_environment(report.get("environment") or {}, labels)
    if env_rows:
        document.add_heading(labels["preview_environment_label"], level=2)
        table = document.add_table(rows=len(env_rows), cols=2)
        table.style = "Light List Accent 1"
        for idx, (name, value) in enumerate(env_rows):
            cell_label = table.rows[idx].cells[0]
            cell_value = table.rows[idx].cells[1]
            cell_label.text = ""
            cell_value.text = ""
            run = cell_label.paragraphs[0].add_run(name)
            run.bold = True
            cell_value.paragraphs[0].add_run(value)

    if report.get("preconditions"):
        document.add_heading(labels["preview_preconditions_label"], level=2)
        for item in report["preconditions"]:
            document.add_paragraph(item, style="List Bullet")

    if report.get("steps_to_reproduce"):
        document.add_heading(labels["preview_str_label"], level=2)
        for step in report["steps_to_reproduce"]:
            document.add_paragraph(step, style="List Number")

    if report.get("expected_result"):
        h_para = document.add_heading(labels["preview_expected_label"], level=2)
        try:
            for run in h_para.runs:
                run.font.color.rgb = RGBColor.from_string("059669")
        except Exception:
            pass
        document.add_paragraph(report["expected_result"])

    if report.get("actual_result"):
        h_para = document.add_heading(labels["preview_actual_label"], level=2)
        try:
            for run in h_para.runs:
                run.font.color.rgb = RGBColor.from_string("B91C1C")
        except Exception:
            pass
        document.add_paragraph(report["actual_result"])

    if images:
        document.add_heading(labels["preview_attachments_label"], level=2)
        for idx, img in enumerate(images, start=1):
            caption = document.add_paragraph()
            caption_run = caption.add_run(f"Screenshot {idx}: {img.name}")
            caption_run.bold = True
            try:
                document.add_picture(img.path, width=Inches(5.5))
            except Exception as exc:
                logger_service.log_exception(
                    "ai_bug_report.pipeline",
                    "write_docx_add_picture_failed",
                    exc,
                    name=img.name,
                )
                document.add_paragraph(
                    f"[Could not embed {img.name}: {exc}]"
                )

    if report.get("attachments_summary"):
        if not images:
            document.add_heading(labels["preview_attachments_label"], level=2)
        for item in report["attachments_summary"]:
            document.add_paragraph(item, style="List Bullet")

    if report.get("additional_notes"):
        document.add_heading(labels["preview_notes_label"], level=2)
        document.add_paragraph(report["additional_notes"])

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))


def save_bug_report_docx(*, labels: dict[str, str]) -> SaveResult:
    """Persist the current ``STATE.last_report`` to disk.

    Writes a fresh run folder via :func:`store.new_run_dir` (so two
    consecutive saves never overwrite each other) and appends a
    ``RunSummary`` entry to ``~/AI Hub/history.json``.
    """
    if not STATE.last_report:
        return SaveResult(ok=False, error="run_first")

    title = (STATE.last_report.get("title") or "").strip() or "Bug report"
    basename = _safe_filename(title)

    _set_activity("saving")
    logger_service.log_event(
        "INFO",
        "ai_bug_report.pipeline",
        "save_bug_report_start",
        title_chars=len(title),
        images=len(STATE.images),
        documents=len(STATE.documents),
    )

    try:
        folder_path = Path(
            store.new_run_dir(
                role=_slug(title),
                company="",
                section="ai_bug_report",
            )
        )
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "save_full_new_run_dir_failed", exc,
        )
        STATE.last_error = str(exc)
        _set_activity("error")
        return SaveResult(ok=False, error=str(exc))

    STATE.last_run_folder = str(folder_path)

    docx_path = folder_path / f"{basename}.docx"
    md_path = folder_path / f"{basename}.md"
    json_path = folder_path / "summary.json"

    try:
        _write_docx(
            docx_path,
            report=STATE.last_report,
            labels=labels,
            images=list(STATE.images),
        )
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "save_bug_report_docx_failed", exc,
            folder=str(folder_path),
        )
        STATE.last_error = str(exc)
        _set_activity("error")
        return SaveResult(ok=False, folder=str(folder_path), error=str(exc))

    try:
        md_path.write_text(
            _render_markdown(STATE.last_report, labels), encoding="utf-8"
        )
    except OSError as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "save_bug_report_md_failed", exc,
            folder=str(folder_path),
        )

    try:
        json_path.write_text(
            json.dumps(
                {
                    "report": STATE.last_report,
                    "saved_at": datetime.now().isoformat(timespec="seconds"),
                    "images": [
                        {"name": img.name, "size_bytes": img.size_bytes}
                        for img in STATE.images
                    ],
                    "documents": [
                        {"name": d.name, "size_bytes": d.size_bytes, "ext": d.ext}
                        for d in STATE.documents
                    ],
                    "cost": {
                        "calls": COST.calls,
                        "tokens": COST.tokens_total,
                        "usd": COST.cost_usd,
                    },
                },
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
    except OSError as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "save_bug_report_json_failed", exc,
            folder=str(folder_path),
        )

    try:
        summary = store.RunSummary(
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M"),
            role=title,
            company="",
            overall_score=0,
            folder=str(folder_path),
            provider="",
            model="",
            cost_usd=float(COST.cost_usd),
            docs=[docx_path.name],
            note="ai_bug_report",
        )
        store.append_run(summary)
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "save_bug_report_history_append_failed", exc,
        )

    STATE.last_save_path = str(docx_path)
    _set_activity("ready")
    logger_service.log_event(
        "INFO",
        "ai_bug_report.pipeline",
        "save_bug_report_done",
        folder=str(folder_path),
        path=str(docx_path),
    )
    return SaveResult(ok=True, folder=str(folder_path), path=str(docx_path))


# ---------------------------------------------------------------------------
# History helpers
# ---------------------------------------------------------------------------


def list_saved_runs() -> list[dict]:
    """Return AI Bug Report runs from disk, newest first.

    Reads ``~/AI Hub/history.json`` via :func:`store.list_runs` and
    keeps only entries whose ``note == "ai_bug_report"`` (so we never
    accidentally show AI Career / AI Jobs runs in our History tab).
    Each entry is enriched with the in-folder ``summary.json`` payload
    when available, so the row can show the bug title, severity, and
    screenshot count instead of just the folder name.
    """
    out: list[dict] = []
    try:
        runs = store.list_runs()
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "list_saved_runs_failed", exc,
        )
        return out

    for run in runs:
        if (run.note or "").strip().lower() != "ai_bug_report":
            continue
        record: dict[str, Any] = {
            "timestamp": run.timestamp,
            "title": run.role,
            "folder": run.folder,
            "severity": "",
            "priority": "",
            "reproducibility": "",
            "image_count": 0,
            "doc_count": 0,
            "docs": list(run.docs or []),
            "cost_usd": float(run.cost_usd or 0.0),
        }
        try:
            summary = store.read_run_summary(run.folder) or {}
            report = summary.get("report") or {}
            if isinstance(report, dict):
                normalised = _normalise_report(report)
                primary = _primary_scenario(normalised)
                if normalised.get("title"):
                    record["title"] = normalised.get("title")
                record["severity"] = primary.get("severity") or ""
                record["priority"] = primary.get("priority") or ""
                record["reproducibility"] = primary.get("reproducibility") or ""
                record["scenario_count"] = len(normalised.get("scenarios") or [])
            images = summary.get("images") or []
            docs = summary.get("documents") or []
            if isinstance(images, list):
                record["image_count"] = len(images)
            if isinstance(docs, list):
                record["doc_count"] = len(docs)
        except Exception as exc:
            logger_service.log_exception(
                "ai_bug_report.pipeline", "list_saved_runs_summary_failed", exc,
                folder=run.folder,
            )
        out.append(record)
    logger_service.log_event(
        "INFO", "ai_bug_report.pipeline", "list_saved_runs_done",
        count=len(out),
    )
    return out


def restore_run(folder: str) -> bool:
    """Load a saved bug-report run back into Preview."""
    if not folder:
        return False
    try:
        summary = store.read_run_summary(folder) or {}
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "restore_run_read_failed", exc, folder=folder,
        )
        STATE.last_error = str(exc)
        return False
    report = summary.get("report")
    if not isinstance(report, dict):
        STATE.last_error = "Saved run does not contain a report JSON."
        logger_service.log_event(
            "WARNING", "ai_bug_report.pipeline", "restore_run_missing_report",
            folder=folder,
        )
        return False
    STATE.last_report = _normalise_report(report)
    STATE.last_run_folder = folder
    docs = summary.get("documents") or []
    if isinstance(docs, list):
        doc_names = [str(item.get("name") or "") for item in docs if isinstance(item, dict)]
    else:
        doc_names = []
    STATE.last_save_path = ""
    logger_service.log_event(
        "INFO", "ai_bug_report.pipeline", "restore_run_done",
        folder=folder,
        scenarios=len(STATE.last_report.get("scenarios") or []),
        documents=len(doc_names),
    )
    REFS.request_context_refresh()
    return True


def delete_run(folder: str) -> bool:
    """Delete a saved run folder + remove it from ``history.json``.

    Returns ``True`` when every disk delete succeeded. Errors are
    logged but never raised - the caller usually surfaces an inline
    toast on failure (see :mod:`src.sections.ai_bug_report.tab_history`).
    """
    if not folder:
        return False
    target = Path(folder)
    success = True
    try:
        if target.exists():
            for child in target.glob("*"):
                try:
                    child.unlink()
                except Exception as exc:
                    logger_service.log_exception(
                        "ai_bug_report.pipeline", "delete_run_unlink_failed", exc,
                        path=str(child),
                    )
                    success = False
            try:
                target.rmdir()
            except Exception as exc:
                logger_service.log_exception(
                    "ai_bug_report.pipeline", "delete_run_rmdir_failed", exc,
                    path=str(target),
                )
                success = False
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "delete_run_failed", exc, folder=folder,
        )
        success = False

    # Drop the matching row from the global ``history.json`` even if
    # the disk delete failed - the user expects the row to disappear
    # from the UI either way (and they will see the error toast).
    try:
        history_path = store.history_path()
        if history_path.exists():
            try:
                raw = json.loads(history_path.read_text(encoding="utf-8")) or []
            except json.JSONDecodeError:
                raw = []
            target_norm = folder.rstrip("\\/").strip()
            keep = [
                row for row in raw
                if (row.get("folder") or "").rstrip("\\/").strip() != target_norm
            ]
            if len(keep) != len(raw):
                history_path.parent.mkdir(parents=True, exist_ok=True)
                history_path.write_text(
                    json.dumps(keep, indent=2, ensure_ascii=False),
                    encoding="utf-8",
                )
    except Exception as exc:
        logger_service.log_exception(
            "ai_bug_report.pipeline", "delete_run_history_rewrite_failed", exc,
            folder=folder,
        )

    STATE.runs_history = [
        entry for entry in STATE.runs_history
        if entry.get("folder") != folder
    ]
    if STATE.last_run_folder == folder:
        STATE.last_run_folder = ""
    logger_service.log_event(
        "INFO", "ai_bug_report.pipeline", "delete_run_done",
        folder=folder, success=success,
    )
    return success


__all__ = [
    "StepResult",
    "SaveResult",
    "load_demo",
    "generate_followup_questions",
    "generate_bug_report",
    "save_bug_report_docx",
    "list_saved_runs",
    "restore_run",
    "delete_run",
]
